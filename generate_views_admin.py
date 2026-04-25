# Script to generate polls/views_admin.py cleanly

with open('polls/views_admin.py', 'w', encoding='utf-8') as f:

    f.write('''from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.http import HttpResponse
from django.db.models import Count
from django.views.decorators.http import require_http_methods
from django.utils import timezone
from django.conf import settings
import csv
import json
import io
import secrets
import string
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import Position, Candidate, PoliceUser, Election, Vote, ElectionPosition, AuditLog
from .forms import PoliceUserRegistrationForm, PoliceUserEditForm, PositionForm, ElectionForm, CandidateForm, BulkVoterUploadForm

def is_admin(user):
    return user.role in ['SUPER_ADMIN', 'ADMIN']

def is_superadmin(user):
    return user.role == 'SUPER_ADMIN'


def create_audit_log(user, action, details, request=None, target_model='', target_id=None):
    ip_address = request.META.get('REMOTE_ADDR') if request else None
    AuditLog.objects.create(
        user=user,
        action=action,
        details=details,
        ip_address=ip_address,
        target_model=target_model,
        target_id=target_id
    )


@login_required
@user_passes_test(is_admin)
def admin_dashboard(request):
    now = timezone.now()
    stats = {
        'positions_count': Position.objects.count(),
        'elections_count': Election.objects.count(),
        'active_elections': Election.objects.filter(start_time__lte=now, end_time__gte=now).count(),
        'candidates_count': Candidate.objects.count(),
        'voters_count': PoliceUser.objects.filter(role='VOTER').count(),
        'is_admin': request.user.role == 'ADMIN',
        'is_superadmin': request.user.role == 'SUPER_ADMIN',
    }
    return render(request, 'polls/admin_dashboard.html', stats)

@login_required
@user_passes_test(is_admin)
def admin_positions(request):
    positions = Position.objects.all().order_by('name')
    if request.method == 'POST':
        if request.user.role != 'SUPER_ADMIN':
            messages.error(request, 'Only Super Admins can create positions.')
            return redirect('polls:admin_positions')
        form = PositionForm(request.POST)
        if form.is_valid():
            position = form.save()
            create_audit_log(
                request.user, AuditLog.ACTION_POSITION_CREATE,
                f"Created position: {position.name}", request,
                target_model='Position', target_id=position.id
            )
            messages.success(request, 'Position created successfully!')
            return redirect('polls:admin_positions')
    else:
        form = PositionForm()
    return render(request, 'polls/admin_positions.html', {'positions': positions, 'form': form})

@login_required
@user_passes_test(is_superadmin)
def edit_position(request, position_id):
    position = get_object_or_404(Position, pk=position_id)
    if request.method == 'POST':
        form = PositionForm(request.POST, instance=position)
        if form.is_valid():
            form.save()
            create_audit_log(
                request.user, AuditLog.ACTION_POSITION_UPDATE,
                f"Updated position: {position.name}", request,
                target_model='Position', target_id=position.id
            )
            messages.success(request, 'Position updated successfully!')
            return redirect('polls:admin_positions')
    else:
        form = PositionForm(instance=position)
    return render(request, 'polls/admin_position_edit.html', {'form': form, 'position': position})

@login_required
@user_passes_test(is_superadmin)
def delete_position(request, position_id):
    position = get_object_or_404(Position, pk=position_id)
    if request.method == 'POST':
        create_audit_log(
            request.user, AuditLog.ACTION_POSITION_DELETE,
            f"Deleted position: {position.name}", request,
            target_model='Position', target_id=position.id
        )
        position.delete()
        messages.success(request, 'Position deleted successfully!')
    return redirect('polls:admin_positions')

@login_required
@user_passes_test(is_admin)
def admin_elections(request):
    elections = Election.objects.all().prefetch_related('positions').order_by('-start_time')
    now = timezone.now()
    positions = Position.objects.all()
    return render(request, 'polls/admin_elections.html', {
        'elections': elections,
        'now': now,
        'positions': positions
    })

@login_required
@user_passes_test(is_superadmin)
def create_election(request):
    from .models import ElectionPosition
    from datetime import datetime
    positions = Position.objects.all()
    
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        start_time = request.POST.get('start_time', '').strip()
        end_time = request.POST.get('end_time', '').strip()
        eligible_ranks = request.POST.get('eligible_ranks', '').strip()
        eligible_stations = request.POST.get('eligible_stations', '').strip()
        selected_positions = request.POST.getlist('positions')
        
        if not title:
            messages.error(request, 'Election title is required.')
            return render(request, 'polls/create_election.html', {
                'positions': positions,
                'selected_position_ids': selected_positions
            })
        
        if not start_time or not end_time:
            messages.error(request, 'Start time and end time are required.')
            return render(request, 'polls/create_election.html', {
                'positions': positions,
                'selected_position_ids': selected_positions
            })
        
        if not selected_positions:
            messages.error(request, 'At least one position must be selected.')
            return render(request, 'polls/create_election.html', {
                'positions': positions,
                'selected_position_ids': selected_positions
            })
        
        try:
            start_dt = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
            end_dt = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')
            start_dt = timezone.make_aware(start_dt)
            end_dt = timezone.make_aware(end_dt)
        except ValueError as e:
            messages.error(request, f'Invalid date/time format. Please use the date/time picker. Error: {str(e)}')
            return render(request, 'polls/create_election.html', {
                'positions': positions,
                'selected_position_ids': selected_positions
            })
        
        if start_dt >= end_dt:
            messages.error(request, 'End time must be after start time.')
            return render(request, 'polls/create_election.html', {
                'positions': positions,
                'selected_position_ids': selected_positions
            })
        
        try:
            election = Election.objects.create(
                title=title,
                description=description,
                start_time=start_dt,
                end_time=end_dt,
                eligible_ranks=eligible_ranks,
                eligible_stations=eligible_stations,
                created_by=request.user
            )
            
            positions_added = 0
            for pos_id in selected_positions:
                try:
                    pos = Position.objects.get(id=pos_id)
                    ElectionPosition.objects.create(election=election, position=pos)
                    positions_added += 1
                except Position.DoesNotExist:
                    pass
            
            if positions_added == 0:
                election.delete()
                messages.error(request, 'No valid positions found. Election not created.')
                return render(request, 'polls/create_election.html', {
                    'positions': positions,
                    'selected_position_ids': selected_positions
                })
            
            create_audit_log(
                request.user, AuditLog.ACTION_ELECTION_CREATE,
                f"Created election: {election.title} (ID: {election.id}, positions: {positions_added})", request,
                target_model='Election', target_id=election.id
            )
            messages.success(request, f'Election "{title}" created successfully with {positions_added} position(s)!')
            return redirect('polls:admin_elections')
        except Exception as e:
            messages.error(request, f'Error creating election: {str(e)}')
            return render(request, 'polls/create_election.html', {
                'positions': positions,
                'selected_position_ids': selected_positions
            })
    
    return render(request, 'polls/create_election.html', {
        'positions': positions,
        'selected_position_ids': []
    })

@login_required
@user_passes_test(is_superadmin)
def edit_election(request, election_id):
    from .models import ElectionPosition
    from datetime import datetime
    
    election = get_object_or_404(Election.objects.prefetch_related('positions'), pk=election_id)
    positions = Position.objects.all()
    selected_position_ids = [str(pos.id) for pos in election.positions.all()]
    
    if request.method == 'POST':
        election.title = request.POST.get('title', election.title)
        election.description = request.POST.get('description', '')
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')
        election.eligible_ranks = request.POST.get('eligible_ranks', '')
        election.eligible_stations = request.POST.get('eligible_stations', '')
        
        if start_time and end_time:
            try:
                start_dt = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
                end_dt = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')
                election.start_time = timezone.make_aware(start_dt)
                election.end_time = timezone.make_aware(end_dt)
            except ValueError as e:
                messages.error(request, f'Invalid date/time format: {str(e)}')
                return render(request, 'polls/create_election.html', {
                    'election': election,
                    'positions': positions,
                    'selected_position_ids': selected_position_ids
                })
        
        election.save()
        
        selected_positions = request.POST.getlist('positions')
        election.positions.clear()
        for pos_id in selected_positions:
            try:
                pos = Position.objects.get(id=pos_id)
                ElectionPosition.objects.create(election=election, position=pos)
            except Position.DoesNotExist:
                pass
        
        create_audit_log(
            request.user, AuditLog.ACTION_ELECTION_UPDATE,
            f"Updated election: {election.title} (ID: {election.id})", request,
            target_model='Election', target_id=election.id
        )
        messages.success(request, 'Election updated successfully!')
        return redirect('polls:admin_elections')
    
    return render(request, 'polls/create_election.html', {
        'election': election,
        'positions': positions,
        'selected_position_ids': selected_position_ids
    })

@login_required
@user_passes_test(is_superadmin)
def toggle_election_status(request, election_id):
    election = get_object_or_404(Election, pk=election_id)
    now_active = election.is_open()
    messages.info(request, f'Election status is automatically determined by start/end times. Currently: {"active" if now_active else "inactive"}.')
    return redirect('polls:admin_elections')

@login_required
@user_passes_test(is_superadmin)
def delete_election(request, election_id):
    election = get_object_or_404(Election, pk=election_id)
    if request.method == 'POST':
        create_audit_log(
            request.user, AuditLog.ACTION_ELECTION_DELETE,
            f"Deleted election: {election.title} (ID: {election.id})", request,
            target_model='Election', target_id=election.id
        )
        election.delete()
        messages.success(request, 'Election deleted successfully!')
    return redirect('polls:admin_elections')

@login_required
@user_passes_test(is_admin)
@require_http_methods(["GET", "POST"])
def admin_register_voter(request):
    import secrets
    import string
    if request.method == 'POST':
        form = PoliceUserRegistrationForm(request.POST)
        if form.is_valid():
            voter = form.save(commit=False)
            auto_password = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(10))
            voter.username = str(voter.force_number)
            voter.role = 'VOTER'
            voter.is_staff = False
            voter.is_active_voter = True
            voter.set_password(auto_password)
            voter.is_active = True
            voter.must_change_password = True
            voter.save()
            create_audit_log(
                request.user, AuditLog.ACTION_VOTER_CREATE,
                f"Registered voter: {voter.username} (Force #{voter.force_number})", request,
                target_model='PoliceUser', target_id=voter.id
            )
            request.session['registered_voter_credentials'] = {
                'username': voter.username,
                'password': auto_password,
                'force_number': voter.force_number,
                'full_name': voter.get_full_name()
            }
            messages.success(request, f'Voter {voter.username} registered successfully!')
            return redirect('polls:admin_voter_credentials')
    else:
        form = PoliceUserRegistrationForm(initial={'role': 'VOTER', 'is_active_voter': True})
        form.fields['role'].disabled = True
    return render(request, 'polls/register_voter.html', {'form': form, 'title': 'Register New Voter'})

@login_required
@user_passes_test(is_admin)
def admin_voter_credentials(request):
    credentials = request.session.pop('registered_voter_credentials', None)
    if not credentials:
        credentials = request.session.pop('reset_voter_credentials', None)
    if not credentials:
        messages.error(request, 'No credentials to display.')
        return redirect('polls:admin_register_voter')
    return render(request, 'polls/voter_credentials.html', {'credentials': credentials})

@login_required
@user_passes_test(is_admin)
def admin_voters(request):
    voters = PoliceUser.objects.filter(role='VOTER').order_by('-date_joined')
    return render(request, 'polls/admin_voters.html', {'voters': voters})

@login_required
@user_passes_test(is_admin)
def admin_edit_voter(request, voter_id):
    voter = get_object_or_404(PoliceUser, pk=voter_id)
    if request.method == 'POST':
        form = PoliceUserEditForm(request.POST, instance=voter)
        if form.is_valid():
            updated_voter = form.save(commit=False)
            updated_voter.role = 'VOTER'
            updated_voter.is_staff = False
            updated_voter.save()
            create_audit_log(
                request.user, AuditLog.ACTION_VOTER_UPDATE,
                f"Updated voter: {voter.username} (Force #{voter.force_number})", request,
                target_model='PoliceUser', target_id=voter.id
            )
            messages.success(request, f'Voter {voter.username} updated successfully!')
            return redirect('polls:admin_voters')
    else:
        form = PoliceUserEditForm(instance=voter)
        form.fields['role'].disabled = True
    return render(request, 'polls/edit_voter.html', {'form': form, 'voter': voter})

@login_required
@user_passes_test(is_admin)
def admin_delete_voter(request, voter_id):
    voter = get_object_or_404(PoliceUser, pk=voter_id)
    if request.method == 'POST':
        create_audit_log(
            request.user, AuditLog.ACTION_VOTER_DELETE,
            f"Deleted voter: {voter.username} (Force #{voter.force_number})", request,
            target_model='PoliceUser', target_id=voter.id
        )
        voter.delete()
        messages.success(request, 'Voter deleted successfully!')
    return redirect('polls:admin_voters')

@login_required
@user_passes_test(is_admin)
def admin_reset_voter_password(request, voter_id):
    import secrets
    import string
    voter = get_object_or_404(PoliceUser, pk=voter_id)
    if request.method == 'POST':
        new_password = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(10))
        voter.set_password(new_password)
        voter.must_change_password = True
        voter.save()
        create_audit_log(
            request.user, AuditLog.ACTION_VOTER_RESET_PASSWORD,
            f"Reset password for voter: {voter.username} (Force #{voter.force_number})", request,
            target_model='PoliceUser', target_id=voter.id
        )
        request.session['reset_voter_credentials'] = {
            'username': voter.username,
            'password': new_password,
            'force_number': voter.force_number,
            'full_name': voter.get_full_name()
        }
        messages.success(request, f'Password reset for {voter.username}!')
        return redirect('polls:admin_voter_credentials')
    return redirect('polls:admin_voters')

@login_required
@user_passes_test(is_superadmin)
def admin_register_candidate(request):
    elections = Election.objects.prefetch_related('positions').all()
    positions = Position.objects.all()
    
    election_positions_map = {}
    for election in elections:
        election_positions_map[str(election.id)] = [
            {'id': pos.id, 'name': pos.name} 
            for pos in election.positions.all()
        ]
    
    if request.method == 'POST':
        form = CandidateForm(request.POST, request.FILES)
        if form.is_valid():
            candidate = form.save(commit=False)
            candidate.created_by = request.user
            candidate.save()
            create_audit_log(
                request.user, AuditLog.ACTION_CANDIDATE_CREATE,
                f"Registered candidate: {candidate.name} (Force #{candidate.force_number}) for {candidate.election.title}", request,
                target_model='Candidate', target_id=candidate.id
            )
            messages.success(request, f'Candidate {candidate.name} registered successfully!')
            return redirect('polls:admin_candidates')
    else:
        form = CandidateForm()
    
    return render(request, 'polls/register_candidate.html', {
        'form': form,
        'elections': elections,
        'positions': positions,
        'election_positions_json': json.dumps(election_positions_map)
    })

@login_required
@user_passes_test(is_superadmin)
def admin_edit_candidate(request, candidate_id):
    candidate = get_object_or_404(Candidate, pk=candidate_id)
    if request.method == 'POST':
        form = CandidateForm(request.POST, request.FILES, instance=candidate)
        if form.is_valid():
            form.save()
            create_audit_log(
                request.user, AuditLog.ACTION_CANDIDATE_UPDATE,
                f"Updated candidate: {candidate.name} (Force #{candidate.force_number})", request,
                target_model='Candidate', target_id=candidate.id
            )
            messages.success(request, f'Candidate {candidate.name} updated successfully!')
            return redirect('polls:admin_candidates')
    else:
        form = CandidateForm(instance=candidate)
    
    elections = Election.objects.prefetch_related('positions').all()
    election_positions_map = {}
    all_positions = Position.objects.all()
    
    for election in elections:
        election_positions = election.positions.all()
        positions_list = [{'id': pos.id, 'name': pos.name} for pos in election_positions]
        
        if not positions_list and candidate.election_id == election.id:
            positions_list.append({
                'id': candidate.position.id,
                'name': f"{candidate.position.name} (Current)"
            })
        
        if not positions_list:
            positions_list = [{'id': pos.id, 'name': pos.name} for pos in all_positions]
        
        election_positions_map[str(election.id)] = positions_list
    
    return render(request, 'polls/edit_candidate.html', {
        'form': form,
        'candidate': candidate,
        'elections': elections,
        'positions': all_positions,
        'election_positions_json': json.dumps(election_positions_map),
        'current_position_id': candidate.position.id
    })

@login_required
@user_passes_test(is_superadmin)
def admin_delete_candidate(request, candidate_id):
    candidate = get_object_or_404(Candidate, pk=candidate_id)
    if request.method == 'POST':
        create_audit_log(
            request.user, AuditLog.ACTION_CANDIDATE_DELETE,
            f"Deleted candidate: {candidate.name} (Force #{candidate.force_number}) from {candidate.election.title}", request,
            target_model='Candidate', target_id=candidate.id
        )
        candidate.delete()
        messages.success(request, 'Candidate deleted successfully!')
    return redirect('polls:admin_candidates')

@login_required
@user_passes_test(is_admin)
def admin_candidates(request):
    candidates = Candidate.objects.all().select_related('election', 'position').order_by('-id')
    return render(request, 'polls/admin_candidates.html', {'candidates': candidates})

@login_required
def view_candidate(request, candidate_id):
    candidate = get_object_or_404(Candidate, pk=candidate_id)
    is_viewer_admin = request.user.role in ['SUPER_ADMIN', 'ADMIN']
    
    if request.user.role == 'VOTER':
        if not candidate.election.is_voter_eligible(request.user):
            messages.error(request, 'You are not eligible to view candidates for this election.')
            return redirect('polls:elections')
    
    votes_count = candidate.vote_set.count()
    election = candidate.election
    total_votes = election.vote_set.count()
    vote_percentage = (votes_count / total_votes * 100) if total_votes > 0 else 0
    
    return render(request, 'polls/view_candidate.html', {
        'candidate': candidate,
        'votes_count': votes_count,
        'vote_percentage': vote_percentage,
        'election': election,
        'is_admin': is_viewer_admin
    })

@login_required
@user_passes_test(is_admin)
def export_voters_csv(request):
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_VOTERS, "Exported voters as CSV", request)
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="voters.csv"'
    writer = csv.writer(response)
    writer.writerow(['Username', 'Force Number', 'Full Name', 'Rank', 'Station', 'Phone', 'Email', 'Role', 'Votes Cast', 'Registered Date', 'Status'])
    voters = PoliceUser.objects.filter(role='VOTER').select_related()
    for voter in voters:
        writer.writerow([
            voter.username, voter.force_number, voter.get_full_name(),
            voter.get_rank_display(), voter.station, voter.phone, voter.email,
            voter.get_role_display(), voter.votes.count(), voter.date_joined.strftime('%Y-%m-%d'),
            'Active' if voter.is_active else 'Inactive'
        ])
    return response

@login_required
@user_passes_test(is_admin)
def export_voters_pdf(request):
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_VOTERS, "Exported voters as PDF", request)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="voters.pdf"'
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    title = Paragraph("Voter Registration Report", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 20))
    
    voters = PoliceUser.objects.filter(role='VOTER').select_related()
    data = [['#', 'Force #', 'Name', 'Rank', 'Station', 'Phone', 'Votes']]
    for i, voter in enumerate(voters, 1):
        data.append([
            str(i), str(voter.force_number), voter.get_full_name(),
            voter.get_rank_display(), voter.station, voter.phone or '-', str(voter.votes.count())
        ])
    
    table = Table(data, colWidths=[0.5*inch, 1*inch, 2*inch, 1.2*inch, 1.5*inch, 1.3*inch, 0.7*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
    ]))
    elements.append(table)
    doc.build(elements)
    return HttpResponse(buffer.getvalue(), content_type='application/pdf')

@login_required
@user_passes_test(is_admin)
def export_voters_docx(request):
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_VOTERS, "Exported voters as DOCX", request)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="voters.docx"'
    doc = Document()
    doc.add_heading('Voter Registration Report', 0)
    doc.add_paragraph(f'Generated: {timezone.now().strftime("%Y-%m-%d %H:%M")}')
    
    voters = PoliceUser.objects.filter(role='VOTER').select_related()
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Force #'
    hdr_cells[2].text = 'Name'
    hdr_cells[3].text = 'Rank'
    hdr_cells[4].text = 'Station'
    hdr_cells[5].text = 'Phone'
    hdr_cells[6].text = 'Votes'
    
    for i, voter in enumerate(voters, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(voter.force_number)
        row_cells[2].text = voter.get_full_name()
        row_cells[3].text = voter.get_rank_display()
        row_cells[4].text = voter.station
        row_cells[5].text = voter.phone or '-'
        row_cells[6].text = str(voter.votes.count())
    
    doc.save(response)
    return response

@login_required
@user_passes_test(is_admin)
def export_candidates_csv(request):
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_CANDIDATES, "Exported candidates as CSV", request)
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="candidates.csv"'
    writer = csv.writer(response)
    writer.writerow(['Name', 'Force Number', 'Rank', 'Position', 'Election', 'Biography', 'Votes'])
    for candidate in Candidate.objects.all().select_related('election', 'position'):
        writer.writerow([
            candidate.name, candidate.force_number, candidate.get_rank_display(),
            candidate.position.name, candidate.election.title,
            candidate.biography[:100], candidate.vote_set.count()
        ])
    return response

@login_required
@user_passes_test(is_admin)
def export_candidates_pdf(request):
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_CANDIDATES, "Exported candidates as PDF", request)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="candidates.pdf"'
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    elements.append(Paragraph("Candidates Report", styles['Title']))
    elements.append(Spacer(1, 20))
    
    candidates = Candidate.objects.all().select_related('election', 'position')
    data = [['Name', 'Position', 'Election', 'Votes']]
    for c in candidates:
        data.append([c.name, c.position.name, c.election.title, str(c.vote_set.count())])
    
    table = Table(data, colWidths=[2*inch, 1.5*inch, 2*inch, 0.8*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    elements.append(table)
    doc.build(elements)
    return HttpResponse(buffer.getvalue(), content_type='application/pdf')

@login_required
@user_passes_test(is_admin)
def export_candidates_docx(request):
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_CANDIDATES, "Exported candidates as DOCX", request)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="candidates.docx"'
    doc = Document()
    doc.add_heading('Candidates Report', 0)
    
    candidates = Candidate.objects.all().select_related('election', 'position')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Position'
    hdr_cells[2].text = 'Election'
    hdr_cells[3].text = 'Votes'
    
    for c in candidates:
        row_cells = table.add_row().cells
        row_cells[0].text = c.name
        row_cells[1].text = c.position.name
        row_cells[2].text = c.election.title
        row_cells[3].text = str(c.vote_set.count())
    
    doc.save(response)
    return response

@login_required
@user_passes_test(is_admin)
def export_results_csv(request, election_id):
    election = get_object_or_404(Election, pk=election_id)
    create_audit_log(
        request.user, AuditLog.ACTION_EXPORT_RESULTS,
        f"Exported results for election: {election.title} (ID: {election.id}) as CSV", request,
        target_model='Election', target_id=election.id
    )
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="results_{election_id}.csv"'
    writer = csv.writer(response)
    writer.writerow(['Candidate', 'Position', 'Rank', 'Force Number', 'Votes', 'Percentage'])
    total = election.vote_set.count()
    for candidate in election.candidates.all():
        votes = candidate.vote_set.count()
        pct = (votes / total * 100) if total > 0 else 0
        writer.writerow([
            candidate.name, candidate.position.name, candidate.get_rank_display(),
            candidate.force_number, votes, f"{pct:.1f}%"
        ])
    return response

@login_required
@user_passes_test(is_admin)
def export_results_pdf(request, election_id):
    election = get_object_or_404(Election, pk=election_id)
    create_audit_log(
        request.user, AuditLog.ACTION_EXPORT_RESULTS,
        f"Exported results for election: {election.title} (ID: {election.id}) as PDF", request,
        target_model='Election', target_id=election.id
    )
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="results_{election_id}.pdf"'
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    elements.append(Paragraph(f"Election Results: {election.title}", styles['Title']))
    elements.append(Paragraph(f"Date: {election.start_time.strftime('%Y-%m-%d')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    total = election.vote_set.count()
    data = [['Candidate', 'Position', 'Votes', '%']]
    for candidate in election.candidates.all():
        votes = candidate.vote_set.count()
        pct = (votes / total * 100) if total > 0 else 0
        data.append([candidate.name, candidate.position.name, str(votes), f"{pct:.1f}%"])
    
    table = Table(data, colWidths=[2.5*inch, 1.5*inch, 1*inch, 0.8*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(table)
    doc.build(elements)
    return HttpResponse(buffer.getvalue(), content_type='application/pdf')

@login_required
@user_passes_test(is_admin)
def export_results_docx(request, election_id):
    election = get_object_or_404(Election, pk=election_id)
    create_audit_log(
        request.user, AuditLog.ACTION_EXPORT_RESULTS,
        f"Exported results for election: {election.title} (ID: {election.id}) as DOCX", request,
        target_model='Election', target_id=election.id
    )
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="results_{election_id}.docx"'
    doc = Document()
    doc.add_heading(f"Election Results: {election.title}", 0)
    doc.add_paragraph(f"Date: {election.start_time.strftime('%Y-%m-%d')}")
    
    total = election.vote_set.count()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Candidate'
    hdr_cells[1].text = 'Position'
    hdr_cells[2].text = 'Votes'
    hdr_cells[3].text = '%'
    
    for candidate in election.candidates.all():
        votes = candidate.vote_set.count()
        pct = (votes / total * 100) if total > 0 else 0
        row_cells = table.add_row().cells
        row_cells[0].text = candidate.name
        row_cells[1].text = candidate.position.name
        row_cells[2].text = str(votes)
        row_cells[3].text = f"{pct:.1f}%"
    
    doc.save(response)
    return response


def _normalize_bulk_key(value):
    return str(value).strip().lower() if value is not None else ''


def _normalize_bulk_value(value):
    if value is None:
        return ''
    return str(value).strip()


def _load_bulk_voter_rows(uploaded_file, extension):
    if
