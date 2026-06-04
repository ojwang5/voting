from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.http import HttpResponse
from django.db.models import Count
from django.views.decorators.http import require_http_methods
from django.utils import timezone
import csv
import json
import io
import secrets
import string
from .models import Position, Candidate, PoliceUser, Election, Vote, ElectionPosition, AuditLog, ElectionRegistration

# Email helpers
from .emails import send_voter_credentials_email, send_election_invitation_email, send_bulk_voter_credentials_email

from .forms import PoliceUserRegistrationForm, PoliceUserEditForm, PositionForm, CandidateForm, BulkVoterUploadForm, AdminChangePasswordForm

from .statistics import get_candidate_pie_stats, render_candidate_pie_chart_png

def is_admin(user):
    return user.role in ['SUPER_ADMIN', 'ADMIN']

def is_superadmin(user):
    return user.role == 'SUPER_ADMIN'

@login_required
@user_passes_test(is_admin)
def election_statistics(request, election_id):
    election = get_object_or_404(Election, pk=election_id)

    position_stats, total_votes = get_candidate_pie_stats(election)

    import base64
    position_charts = {}
    all_stats = []
    for position, pos_stats in position_stats.items():
        all_stats.extend(pos_stats)
        chart_png = render_candidate_pie_chart_png(pos_stats)
        if chart_png:
            position_charts[position] = "data:image/png;base64," + base64.b64encode(chart_png).decode("ascii")

    return render(
        request,
        "polls/election_statistics.html",
        {
            "election": election,
            "stats": all_stats,
            "position_stats": position_stats,
            "position_charts": position_charts,
            "total_votes": total_votes,
        },
    )

@login_required
@user_passes_test(is_admin)
def export_election_stats_pdf(request, election_id):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    election = get_object_or_404(Election, pk=election_id)
    create_audit_log(
        request.user,
        AuditLog.ACTION_EXPORT_RESULTS,
        f"Exported election statistics for {election.title} as PDF",
        request,
        "Election",
        election.id,
    )

    position_stats, total_votes = get_candidate_pie_stats(election)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=18)
    elements = []
    styles = getSampleStyleSheet()

    if election.logo:
        try:
            from reportlab.platypus import Image
            logo_path = election.logo.path
            logo = Image(logo_path, width=1.5 * inch, height=0.75 * inch)
            logo.hAlign = "CENTER"
            elements.append(logo)
            elements.append(Spacer(1, 12))
        except:
            pass

    elements.append(Paragraph(f"Election Statistics: {election.title}", styles["Title"]))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(f"Total votes cast: {total_votes}", styles["BodyText"]))
    elements.append(Spacer(1, 12))

    from reportlab.platypus import Image as RLImage
    from reportlab.lib.utils import ImageReader

    for position, pos_stats in position_stats.items():
        elements.append(Paragraph(position.name, styles["Heading2"]))
        chart_png = render_candidate_pie_chart_png(pos_stats)
        if chart_png:
            img = ImageReader(io.BytesIO(chart_png))
            elements.append(RLImage(img, width=360, height=360))
            elements.append(Spacer(1, 8))

        data = [["Candidate", "Force #", "Votes", "Share"]]
        for s in pos_stats:
            try:
                pct_val = float(getattr(s, "percentage", 0.0))
            except (TypeError, ValueError):
                pct_val = 0.0
            data.append([s.candidate_name, str(s.force_number), str(s.votes), f"{pct_val:.1f}%"])

        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 10),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 16))

    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = f"attachment; filename=statistics_{election_id}.pdf"
    response.write(pdf)
    return response

@login_required
@user_passes_test(is_admin)
def export_election_stats_docx(request, election_id):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    election = get_object_or_404(Election, pk=election_id)
    create_audit_log(
        request.user,
        AuditLog.ACTION_EXPORT_RESULTS,
        f"Exported election statistics for {election.title} as DOCX",
        request,
        "Election",
        election.id,
    )

    position_stats, total_votes = get_candidate_pie_stats(election)

    doc = Document()

    if election.logo:
        try:
            doc.add_picture(election.logo.path, width=Inches(2.5))
        except:
            pass

    title = doc.add_heading(f"Election Statistics: {election.title}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Total votes cast: {total_votes}")
    doc.add_paragraph("")

    for position, pos_stats in position_stats.items():
        doc.add_heading(position.name, level=2)
        chart_png = render_candidate_pie_chart_png(pos_stats)
        if chart_png:
            try:
                tmp = io.BytesIO(chart_png)
                doc.add_picture(tmp, width=Inches(4.0))
            except:
                pass

        doc.add_paragraph("")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        headers = ["Candidate", "Force #", "Votes", "Share"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                for run in p.runs:
                    run.font.bold = True

        for s in pos_stats:
            row_cells = table.add_row().cells
            row_cells[0].text = s.candidate_name
            row_cells[1].text = str(s.force_number)
            row_cells[2].text = str(s.votes)
            row_cells[3].text = f"{s.percentage:.1f}%"

        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f"attachment; filename=statistics_{election_id}.docx"
    return response

def auto_assign_elections(voter, registered_by=None):
    for election in Election.objects.all():
        if election.is_voter_eligible(voter):
            ElectionRegistration.objects.get_or_create(
                voter=voter, election=election,
                defaults={'registered_by': registered_by}
            )

def create_audit_log(user, action, details, request=None, target_model='', target_id=None):
    ip_address = request.META.get('REMOTE_ADDR') if request else None
    AuditLog.objects.create(
        user=user, action=action, details=details,
        ip_address=ip_address, target_model=target_model, target_id=target_id
    )

@login_required
@user_passes_test(is_admin)
def admin_dashboard(request):
    now = timezone.now()
    
    # Get elections by status
    all_elections = Election.objects.prefetch_related('positions', 'candidates').all()
    active_elections_list = []
    upcoming_elections = []
    ended_elections_list = []
    
    for election in all_elections:
        status = election.status
        if status == 'ACTIVE':
            active_elections_list.append(election)
        elif status == 'UPCOMING':
            upcoming_elections.append(election)
        else:  # ENDED
            ended_elections_list.append(election)
    
    return render(request, 'polls/admin_dashboard.html', {
        'positions_count': Position.objects.count(),
        'elections_count': Election.objects.count(),
        'active_elections': len(active_elections_list),
        'candidates_count': Candidate.objects.count(),
        'voters_count': PoliceUser.objects.filter(role='VOTER').count(),
        'is_admin': request.user.role == 'ADMIN',
        'is_superadmin': request.user.role == 'SUPER_ADMIN',
        'active_elections_list': active_elections_list,
        'upcoming_elections': upcoming_elections,
        'ended_elections_list': ended_elections_list,
    })

@login_required
@user_passes_test(is_admin)
def admin_positions(request):
    positions = Position.objects.all().order_by('name')
    if request.method == 'POST':
        if request.user.role != 'SUPER_ADMIN':
            messages.error(request, 'Only Super Admins can create positions.')
            return redirect('polls:admin_positions')
        form = PositionForm(request.POST)
        if form.is_valid():
            position = form.save()
            create_audit_log(request.user, AuditLog.ACTION_POSITION_CREATE,
                f'Created position: {position.name}', request, 'Position', position.id)
            messages.success(request, 'Position created successfully!')
            return redirect('polls:admin_positions')
    else:
        form = PositionForm()
    return render(request, 'polls/admin_positions.html', {
        'positions': positions, 'form': form,
        'is_superadmin': request.user.role == 'SUPER_ADMIN'
    })

@login_required
@user_passes_test(is_superadmin)
def edit_position(request, position_id):
    position = get_object_or_404(Position, pk=position_id)
    if request.method == 'POST':
        form = PositionForm(request.POST, instance=position)
        if form.is_valid():
            form.save()
            create_audit_log(request.user, AuditLog.ACTION_POSITION_UPDATE,
                f'Updated position: {position.name}', request, 'Position', position.id)
@login_required
@user_passes_test(is_superadmin)
def delete_position(request, position_id):
    position = get_object_or_404(Position, pk=position_id)
    if request.method == 'POST':
        create_audit_log(request.user, AuditLog.ACTION_POSITION_DELETE,
            f'Deleted position: {position.name}', request, 'Position', position.id)
@login_required
@user_passes_test(is_admin)
def admin_elections(request):
    elections = Election.objects.all().prefetch_related('positions').order_by('-start_time')
    return render(request, 'polls/admin_elections.html', {
        'elections': elections,
        'now': timezone.now(),
        'positions': Position.objects.all(),
        'is_superadmin': request.user.role == 'SUPER_ADMIN'
    })

@login_required
@user_passes_test(is_superadmin)
def create_election(request):
    from .models import ElectionPosition
    from datetime import datetime
    positions = Position.objects.all()
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        start_time = request.POST.get('start_time', '').strip()
        end_time = request.POST.get('end_time', '').strip()
        eligible_ranks = request.POST.get('eligible_ranks', '').strip()
        eligible_stations = request.POST.get('eligible_stations', '').strip()
        logo_file = request.FILES.get('logo')
        selected_positions = request.POST.getlist('positions')
        if not title:
            messages.error(request, 'Election title is required.')
            return render(request, 'polls/create_election.html', {'positions': positions, 'selected_position_ids': selected_positions})
        if not start_time or not end_time:
            messages.error(request, 'Start time and end time are required.')
            return render(request, 'polls/create_election.html', {'positions': positions, 'selected_position_ids': selected_positions})
        if not selected_positions:
            messages.error(request, 'At least one position must be selected.')
            return render(request, 'polls/create_election.html', {'positions': positions, 'selected_position_ids': selected_positions})
        try:
            start_dt = timezone.make_aware(datetime.strptime(start_time, '%Y-%m-%dT%H:%M'))
            end_dt = timezone.make_aware(datetime.strptime(end_time, '%Y-%m-%dT%H:%M'))
        except ValueError as e:
            messages.error(request, f'Invalid date/time format: {str(e)}')
            return render(request, 'polls/create_election.html', {'positions': positions, 'selected_position_ids': selected_positions})
        if start_dt >= end_dt:
            messages.error(request, 'End time must be after start time.')
            return render(request, 'polls/create_election.html', {'positions': positions, 'selected_position_ids': selected_positions})
        try:
            election = Election.objects.create(title=title, description=description,
                logo=logo_file if logo_file else None,
                start_time=start_dt, end_time=end_dt, eligible_ranks=eligible_ranks,
                eligible_stations=eligible_stations, created_by=request.user)
            positions_added = 0
            for pos_id in selected_positions:
                try:
                    pos = Position.objects.get(id=pos_id)
                    ElectionPosition.objects.create(election=election, position=pos)
                    positions_added += 1
                except Position.DoesNotExist:
                    pass
            if positions_added == 0:
                election.delete()
                messages.error(request, 'No valid positions found. Election not created.')
                return render(request, 'polls/create_election.html', {'positions': positions, 'selected_position_ids': selected_positions})
            create_audit_log(request.user, AuditLog.ACTION_ELECTION_CREATE,
                f'Created election: {election.title} (ID: {election.id}, positions: {positions_added})', request, 'Election', election.id)
            messages.success(request, f'Election "{title}" created with {positions_added} position(s)!')
            return redirect('polls:admin_elections')
        except Exception as e:
            messages.error(request, f'Error creating election: {str(e)}')
            return render(request, 'polls/create_election.html', {'positions': positions, 'selected_position_ids': selected_positions})
    return render(request, 'polls/create_election.html', {'positions': positions, 'selected_position_ids': []})

@login_required
@user_passes_test(is_superadmin)
def edit_election(request, election_id):
    from .models import ElectionPosition
    from datetime import datetime
    election = get_object_or_404(Election.objects.prefetch_related('positions'), pk=election_id)
    positions = Position.objects.all()
    selected_position_ids = [str(pos.id) for pos in election.positions.all()]
    if request.method == 'POST':
        election.title = request.POST.get('title', election.title)
        election.description = request.POST.get('description', '')
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')
        election.eligible_ranks = request.POST.get('eligible_ranks', '')
        election.eligible_stations = request.POST.get('eligible_stations', '')
        logo_file = request.FILES.get('logo')
        if start_time and end_time:
            try:
                election.start_time = timezone.make_aware(datetime.strptime(start_time, '%Y-%m-%dT%H:%M'))
                election.end_time = timezone.make_aware(datetime.strptime(end_time, '%Y-%m-%dT%H:%M'))
            except ValueError as e:
                messages.error(request, f'Invalid date/time format: {str(e)}')
                return render(request, 'polls/create_election.html', {'election': election, 'positions': positions, 'selected_position_ids': selected_position_ids})
        election.save()
        if logo_file:
            election.logo = logo_file
            election.save()
        selected_positions = request.POST.getlist('positions')
        election.positions.clear()
        for pos_id in selected_positions:
            try:
                pos = Position.objects.get(id=pos_id)
                ElectionPosition.objects.create(election=election, position=pos)
            except Position.DoesNotExist:
                pass
        audit_details = f'Updated election: {election.title} (ID: {election.id})'
        if logo_file:
            audit_details += ' with new logo'
        create_audit_log(request.user, AuditLog.ACTION_ELECTION_UPDATE,
            audit_details, request, 'Election', election.id)
        messages.success(request, 'Election updated successfully!')
        return redirect('polls:admin_elections')
    return render(request, 'polls/create_election.html', {'election': election, 'positions': positions, 'selected_position_ids': selected_position_ids})

@login_required
@user_passes_test(is_superadmin)
def toggle_election_status(request, election_id):
    election = get_object_or_404(Election, pk=election_id)
    messages.info(request, f'Election status is automatically determined by start/end times. Currently: {"active" if election.is_open() else "inactive"}.')
    return redirect('polls:admin_elections')

@login_required
@user_passes_test(is_superadmin)
def delete_election(request, election_id):
    election = get_object_or_404(Election, pk=election_id)
    if request.method == 'POST':
        create_audit_log(request.user, AuditLog.ACTION_ELECTION_DELETE,
            f'Deleted election: {election.title} (ID: {election.id})', request, 'Election', election.id)
        election.delete()
        messages.success(request, 'Election deleted successfully!')
    return redirect('polls:admin_elections')

@login_required
@user_passes_test(is_admin)
@require_http_methods(['GET', 'POST'])
def admin_register_voter(request):
    if request.method == 'POST':
        form = PoliceUserRegistrationForm(request.POST)
        if form.is_valid():
            voter = form.save(commit=False)
            auto_password = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(10))
            voter.username = str(voter.force_number)
            voter.role = 'VOTER'
            voter.is_staff = False
            voter.is_active_voter = True
            voter.set_password(auto_password)
            voter.is_active = True
            voter.must_change_password = True
            voter.save()
            create_audit_log(request.user, AuditLog.ACTION_VOTER_CREATE,
                f'Registered voter: {voter.username} (Force #{voter.force_number})', request, 'PoliceUser', voter.id)
            
            # Register voter for selected elections
            assigned_elections = []
            selected_election_ids = request.POST.getlist('elections')
            if selected_election_ids:
                for eid in selected_election_ids:
                    try:
                        election = Election.objects.get(id=eid)
                        ElectionRegistration.objects.get_or_create(
                            voter=voter, election=election,
                            defaults={'registered_by': request.user}
                        )
                        assigned_elections.append(election.title)
                    except Election.DoesNotExist:
                        pass
            else:
                auto_assign_elections(voter, registered_by=request.user)
                assigned_count = ElectionRegistration.objects.filter(voter=voter).count()
                if assigned_count:
                    assigned_elections = [f'{assigned_count} election(s)']
            
            # Store credentials and election info in session
            request.session['registered_voter_credentials'] = {
                'username': voter.username, 'password': auto_password,
                'force_number': voter.force_number, 'full_name': voter.get_full_name(),
                'assigned_elections': ', '.join(assigned_elections) if assigned_elections else None
            }
            
            # Attempt to send credentials email
            email_sent = send_voter_credentials_email(voter, auto_password)
            if email_sent:
                messages.success(request, f'Voter {voter.username} registered successfully! Credentials sent to {voter.email}')
            else:
                messages.success(request, f'Voter {voter.username} registered successfully!')
                if voter.email:
                    messages.warning(request, f'Could not send email to {voter.email}. Displaying credentials on screen.')
                else:
                    messages.warning(request, 'No email address on file. Displaying credentials on screen.')
            
            return redirect('polls:admin_voter_credentials')
    else:
        form = PoliceUserRegistrationForm(initial={'role': 'VOTER', 'is_active_voter': True})
        form.fields['role'].disabled = True
    upcoming_elections = Election.objects.filter(start_time__gt=timezone.now()).prefetch_related('positions').order_by('start_time')
    return render(request, 'polls/register_voter.html', {
        'form': form, 'title': 'Register New Voter',
        'upcoming_elections': upcoming_elections
    })

@login_required
@user_passes_test(is_admin)
def admin_voter_credentials(request):
    credentials = request.session.pop('registered_voter_credentials', None)
    if not credentials:
        credentials = request.session.pop('reset_voter_credentials', None)
    if not credentials:
        messages.error(request, 'No credentials to display.')
        return redirect('polls:admin_register_voter')
    return render(request, 'polls/voter_credentials.html', {'credentials': credentials})


@login_required
@user_passes_test(is_admin)
def admin_voters(request):
    voters = PoliceUser.objects.filter(role='VOTER').order_by('-date_joined')
    return render(request, 'polls/admin_voters.html', {
        'voters': voters, 'is_superadmin': request.user.role == 'SUPER_ADMIN'
    })

@login_required
@user_passes_test(is_admin)
def admin_edit_voter(request, voter_id):
    voter = get_object_or_404(PoliceUser, pk=voter_id)
    if request.method == 'POST':
        form = PoliceUserEditForm(request.POST, instance=voter)
        if form.is_valid():
            updated_voter = form.save(commit=False)
            updated_voter.role = 'VOTER'
            updated_voter.is_staff = False
            updated_voter.save()
            create_audit_log(request.user, AuditLog.ACTION_VOTER_UPDATE,
                f'Updated voter: {voter.username} (Force #{voter.force_number})', request, 'PoliceUser', voter.id)
            messages.success(request, f'Voter {voter.username} updated successfully!')
            return redirect('polls:admin_voters')
    else:
        form = PoliceUserEditForm(instance=voter)
        form.fields['role'].disabled = True
    return render(request, 'polls/edit_voter.html', {'form': form, 'voter': voter})

@login_required
@user_passes_test(is_admin)
def admin_delete_voter(request, voter_id):
    voter = get_object_or_404(PoliceUser, pk=voter_id)
    if request.method == 'POST':
        create_audit_log(request.user, AuditLog.ACTION_VOTER_DELETE,
            f'Deleted voter: {voter.username} (Force #{voter.force_number})', request, 'PoliceUser', voter.id)
        voter.delete()
        messages.success(request, 'Voter deleted successfully!')
    return redirect('polls:admin_voters')


@login_required
@user_passes_test(is_admin)
def admin_reset_voter_password(request, voter_id):
    import secrets
    import string
    voter = get_object_or_404(PoliceUser, pk=voter_id)
    if request.method == 'POST':
        new_password = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(10))
        voter.set_password(new_password)
        voter.must_change_password = True
        voter.save()
        create_audit_log(request.user, AuditLog.ACTION_VOTER_RESET_PASSWORD,
            f'Reset password for voter: {voter.username} (Force #{voter.force_number})', request, 'PoliceUser', voter.id)
        request.session['reset_voter_credentials'] = {
            'username': voter.username, 'password': new_password,
            'force_number': voter.force_number, 'full_name': voter.get_full_name()
        }
        messages.success(request, f'Password reset for {voter.username}!')
        return redirect('polls:admin_voter_credentials')
    return redirect('polls:admin_voters')

@login_required
@user_passes_test(is_superadmin)
def admin_change_voter_password(request, voter_id):
    voter = get_object_or_404(PoliceUser, pk=voter_id)
    if request.method == 'POST':
        form = AdminChangePasswordForm(request.POST)
        if form.is_valid():
            voter.set_password(form.cleaned_data['new_password'])
            voter.must_change_password = form.cleaned_data.get('force_change', False)
            voter.save()
            create_audit_log(request.user, AuditLog.ACTION_VOTER_RESET_PASSWORD,
                f'Changed password for voter: {voter.username} (Force #{voter.force_number})', request, 'PoliceUser', voter.id)
            messages.success(request, f'Password changed for {voter.username}!')
            return redirect('polls:admin_voters')
    else:
        form = AdminChangePasswordForm()
    return render(request, 'polls/admin_change_password.html', {
        'form': form, 'voter': voter, 'title': 'Change Voter Password'
    })

@login_required
@user_passes_test(is_admin)
def admin_voter_elections(request, voter_id):
    voter = get_object_or_404(PoliceUser, pk=voter_id, role='VOTER')
    registered_election_ids = set(ElectionRegistration.objects.filter(
        voter=voter
    ).values_list('election_id', flat=True))

    registered_elections = Election.objects.prefetch_related('positions').filter(
        id__in=registered_election_ids
    ).order_by('-start_time')

    eligible_unregistered = []
    for election in Election.objects.prefetch_related('positions').all().order_by('-start_time'):
        if election.is_voter_eligible(voter) and election.id not in registered_election_ids:
            eligible_unregistered.append(election)

    return render(request, 'polls/admin_voter_elections.html', {
        'voter': voter,
        'registered_elections': registered_elections,
        'eligible_unregistered': eligible_unregistered,
        'registered_count': registered_elections.count(),
        'eligible_count': len(eligible_unregistered),
        'now': timezone.now(),
    })

@login_required
@user_passes_test(is_admin)
def admin_register_voter_to_election_by_voter(request, voter_id, election_id):
    election = get_object_or_404(Election, pk=election_id)
    voter = get_object_or_404(PoliceUser, pk=voter_id, role='VOTER')

    if not election.is_voter_eligible(voter):
        messages.error(request, f'Voter {voter.get_full_name()} is not eligible for this election.')
        return redirect('polls:admin_voter_elections', voter_id=voter_id)

    if ElectionRegistration.objects.filter(voter=voter, election=election).exists():
        messages.info(request, f'Voter {voter.get_full_name()} is already registered for this election.')
        return redirect('polls:admin_voter_elections', voter_id=voter_id)

    ElectionRegistration.objects.create(voter=voter, election=election, registered_by=request.user)
    create_audit_log(
        request.user, AuditLog.ACTION_VOTER_UPDATE,
        f'Registered voter {voter.username} for election: {election.title}', request,
        target_model='Election', target_id=election.id
    )
    messages.success(request, f'Voter {voter.get_full_name()} registered for "{election.title}".')
    return redirect('polls:admin_voter_elections', voter_id=voter_id)

@login_required
@user_passes_test(is_admin)
def admin_unregister_voter_from_election_by_voter(request, voter_id, election_id):
    election = get_object_or_404(Election, pk=election_id)
    voter = get_object_or_404(PoliceUser, pk=voter_id, role='VOTER')

    registration = ElectionRegistration.objects.filter(voter=voter, election=election).first()
    if not registration:
        messages.error(request, f'Voter {voter.get_full_name()} is not registered for this election.')
        return redirect('polls:admin_voter_elections', voter_id=voter_id)

    if Vote.objects.filter(voter=voter, election=election).exists():
        messages.error(request, f'Cannot unregister {voter.get_full_name()} — they have already voted in this election.')
        return redirect('polls:admin_voter_elections', voter_id=voter_id)

    registration.delete()
    create_audit_log(
        request.user, AuditLog.ACTION_VOTER_UPDATE,
        f'Unregistered voter {voter.username} from election: {election.title}', request,
        target_model='Election', target_id=election.id
    )
    messages.success(request, f'Voter {voter.get_full_name()} unregistered from "{election.title}".')
    return redirect('polls:admin_voter_elections', voter_id=voter_id)

@login_required
@user_passes_test(is_superadmin)
def admin_register_candidate(request):
    elections = Election.objects.prefetch_related('positions').all()
    positions = Position.objects.all()
    election_positions_map = {}
    for election in elections:
        election_positions_map[str(election.id)] = [{'id': pos.id, 'name': pos.name} for pos in election.positions.all()]
    if request.method == 'POST':
        form = CandidateForm(request.POST, request.FILES)
        if form.is_valid():
            candidate = form.save(commit=False)
            candidate.created_by = request.user
            candidate.save()
            create_audit_log(request.user, AuditLog.ACTION_CANDIDATE_CREATE,
                f'Registered candidate: {candidate.name} (Force #{candidate.force_number}) for {candidate.election.title}',
                request, 'Candidate', candidate.id)
            messages.success(request, f'Candidate {candidate.name} registered successfully!')
            return redirect('polls:admin_candidates')
    else:
        form = CandidateForm()
    return render(request, 'polls/register_candidate.html', {
        'form': form, 'elections': elections, 'positions': positions,
        'election_positions_json': json.dumps(election_positions_map)
    })


@login_required
@user_passes_test(is_superadmin)
def admin_edit_candidate(request, candidate_id):
    candidate = get_object_or_404(Candidate, pk=candidate_id)
    if request.method == 'POST':
        form = CandidateForm(request.POST, request.FILES, instance=candidate)
        if form.is_valid():
            form.save()
            create_audit_log(request.user, AuditLog.ACTION_CANDIDATE_UPDATE,
                f'Updated candidate: {candidate.name} (Force #{candidate.force_number})', request, 'Candidate', candidate.id)
            messages.success(request, f'Candidate {candidate.name} updated successfully!')
            return redirect('polls:admin_candidates')
    else:
        form = CandidateForm(instance=candidate)
    elections = Election.objects.prefetch_related('positions').all()
    election_positions_map = {}
    all_positions = Position.objects.all()
    for election in elections:
        election_positions = election.positions.all()
        positions_list = [{'id': pos.id, 'name': pos.name} for pos in election_positions]
        if not positions_list and candidate.election_id == election.id:
            positions_list.append({'id': candidate.position.id, 'name': f'{candidate.position.name} (Current)'})
        if not positions_list:
            positions_list = [{'id': pos.id, 'name': pos.name} for pos in all_positions]
        election_positions_map[str(election.id)] = positions_list
    return render(request, 'polls/edit_candidate.html', {
        'form': form, 'candidate': candidate, 'elections': elections,
        'positions': all_positions, 'election_positions_json': json.dumps(election_positions_map),
        'current_position_id': candidate.position.id
    })


@login_required
@user_passes_test(is_superadmin)
def admin_delete_candidate(request, candidate_id):
    candidate = get_object_or_404(Candidate, pk=candidate_id)
    if request.method == 'POST':
        election_title = candidate.election.title
        create_audit_log(request.user, AuditLog.ACTION_CANDIDATE_DELETE,
            f'Deleted candidate: {candidate.name} (Force #{candidate.force_number}) from {election_title}',
            request, 'Candidate', candidate.id)
        candidate.delete()
        messages.success(request, 'Candidate deleted successfully!')
    return redirect('polls:admin_candidates')

@login_required
@user_passes_test(is_superadmin)
def delete_election_position(request, election_id, position_id):
    election = get_object_or_404(Election, pk=election_id)
    position = get_object_or_404(Position, pk=position_id)
    
    # Safety check: cannot delete if candidates exist for this position in this election
    if Candidate.objects.filter(election=election, position=position).exists():
        messages.error(request, f'Cannot delete position "{position.name}" from "{election.title}". Candidates exist for this position.')
        return redirect('polls:admin_elections')
    
    # Delete the association
    election_pos = ElectionPosition.objects.filter(election=election, position=position)
    count = election_pos.count()
    election_pos.delete()
    
    create_audit_log(request.user, AuditLog.ACTION_ELECTION_UPDATE,
        f'Removed position "{position.name}" from election "{election.title}" ({count} association(s) deleted)',
        request, 'ElectionPosition', position.id)
    
    messages.success(request, f'Position "{position.name}" removed from "{election.title}".')
    return redirect('polls:admin_elections')


@login_required
@user_passes_test(is_admin)
def admin_candidates(request):
    candidates = Candidate.objects.all().select_related('election', 'position').order_by('-id')
    return render(request, 'polls/admin_candidates.html', {
        'candidates': candidates, 'is_superadmin': request.user.role == 'SUPER_ADMIN'
    })


@login_required
def view_candidate(request, candidate_id):
    candidate = get_object_or_404(Candidate, pk=candidate_id)
    if request.user.role == 'VOTER':
        if not candidate.election.is_voter_eligible(request.user):
            messages.error(request, 'You are not eligible to view candidates for this election.')
            return redirect('polls:elections')
    votes_count = candidate.vote_set.count()
    election = candidate.election
    total_votes = election.vote_set.count()
    vote_percentage = (votes_count / total_votes * 100) if total_votes > 0 else 0
    return render(request, 'polls/view_candidate.html', {
        'candidate': candidate, 'votes_count': votes_count,
        'vote_percentage': vote_percentage, 'election': election,
        'is_admin': request.user.role in ['SUPER_ADMIN', 'ADMIN']
    })


@login_required
@user_passes_test(is_admin)
def export_voters_csv(request):
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_VOTERS, 'Exported voters as CSV', request)
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="voters.csv"'
    writer = csv.writer(response)
    writer.writerow(['Username', 'Force Number', 'Full Name', 'Rank', 'Station', 'Phone', 'Email', 'Votes Cast', 'Registered Date', 'Status'])
    for voter in PoliceUser.objects.filter(role='VOTER'):
        writer.writerow([voter.username, voter.force_number, voter.get_full_name(),
            voter.get_rank_display(), voter.station, voter.phone, voter.email,
            voter.votes.count(), voter.date_joined.strftime('%Y-%m-%d'),
            'Active' if voter.is_active else 'Inactive'])
    return response


@login_required
@user_passes_test(is_admin)
def export_voters_pdf(request):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_VOTERS, 'Exported voters as PDF', request)
    voters = PoliceUser.objects.filter(role='VOTER').order_by('force_number')
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=18)
    elements = []
    styles = getSampleStyleSheet()
    elements.append(Paragraph('Voters Export', styles['Title']))
    elements.append(Spacer(1, 12))
    data = [['Username', 'Force #', 'Full Name', 'Rank', 'Station', 'Phone', 'Email', 'Votes', 'Status']]
    for v in voters:
        data.append([v.username, v.force_number, v.get_full_name(), v.get_rank_display(), v.station, v.phone or '-', v.email or '-', str(v.votes.count()), 'Active' if v.is_active else 'Inactive'])
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
    ]))
    elements.append(t)
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="voters.pdf"'
    response.write(pdf)
    return response


@login_required
@user_passes_test(is_admin)
def export_voters_docx(request):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_VOTERS, 'Exported voters as DOCX', request)
    voters = PoliceUser.objects.filter(role='VOTER').order_by('force_number')
    doc = Document()
    title = doc.add_heading('Voters Export', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Username', 'Force #', 'Full Name', 'Rank', 'Station', 'Phone', 'Email', 'Votes', 'Status']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for v in voters:
        row_cells = table.add_row().cells
        row_cells[0].text = v.username
        row_cells[1].text = str(v.force_number)
        row_cells[2].text = v.get_full_name()
        row_cells[3].text = v.get_rank_display()
        row_cells[4].text = v.station
        row_cells[5].text = v.phone or '-'
        row_cells[6].text = v.email or '-'
        row_cells[7].text = str(v.votes.count())
        row_cells[8].text = 'Active' if v.is_active else 'Inactive'
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="voters.docx"'
    return response

@login_required
@user_passes_test(is_admin)
def export_candidates_csv(request):
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_CANDIDATES, 'Exported candidates as CSV', request)
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="candidates.csv"'
    writer = csv.writer(response)
    writer.writerow(['Name', 'Force Number', 'Rank', 'Position', 'Election', 'Votes'])
    for candidate in Candidate.objects.all().select_related('election', 'position'):
        writer.writerow([candidate.name, candidate.force_number, candidate.get_rank_display(),
            candidate.position.name, candidate.election.title, candidate.vote_set.count()])
    return response

@login_required
@user_passes_test(is_admin)
def export_candidates_pdf(request):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_CANDIDATES, 'Exported candidates as PDF', request)
    candidates = Candidate.objects.all().select_related('election', 'position').order_by('position__name', 'name')
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=18)
    elements = []
    styles = getSampleStyleSheet()
    elements.append(Paragraph('Candidates Export', styles['Title']))
    elements.append(Spacer(1, 12))
    data = [['Name', 'Force #', 'Rank', 'Position', 'Election', 'Votes']]
    for c in candidates:
        data.append([c.name, c.force_number, c.get_rank_display(), c.position.name, c.election.title, str(c.vote_set.count())])
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
    ]))
    elements.append(t)
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="candidates.pdf"'
    response.write(pdf)
    return response

@login_required
@user_passes_test(is_admin)
def export_candidates_docx(request):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_CANDIDATES, 'Exported candidates as DOCX', request)
    candidates = Candidate.objects.all().select_related('election', 'position').order_by('position__name', 'name')
    doc = Document()
    title = doc.add_heading('Candidates Export', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Name', 'Force #', 'Rank', 'Position', 'Election', 'Votes']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for c in candidates:
        row_cells = table.add_row().cells
        row_cells[0].text = c.name
        row_cells[1].text = str(c.force_number)
        row_cells[2].text = c.get_rank_display()
        row_cells[3].text = c.position.name
        row_cells[4].text = c.election.title
        row_cells[5].text = str(c.vote_set.count())
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="candidates.docx"'
    return response

@login_required
@user_passes_test(is_admin)
def export_results_csv(request, election_id):
    election = get_object_or_404(Election, pk=election_id)
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_RESULTS,
        f'Exported results for election: {election.title} (ID: {election.id}) as CSV', request, 'Election', election.id)
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="results_{election_id}.csv"'
    writer = csv.writer(response)
    writer.writerow(['Candidate', 'Position', 'Rank', 'Force Number', 'Votes', 'Percentage'])
    total = election.vote_set.count()
    for candidate in election.candidates.all():
        votes = candidate.vote_set.count()
        pct = (votes / total * 100) if total > 0 else 0
        writer.writerow([candidate.name, candidate.position.name, candidate.get_rank_display(),
            candidate.force_number, votes, f'{pct:.1f}%'])
    return response

@login_required
@user_passes_test(is_admin)
def export_results_pdf(request, election_id):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    election = get_object_or_404(Election.objects.prefetch_related('positions'), pk=election_id)
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_RESULTS,
        f'Exported results for election: {election.title} (ID: {election.id}) as PDF', request, 'Election', election.id)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=18)
    elements = []
    styles = getSampleStyleSheet()
    # Add logo if available
    if election.logo:
        try:
            from reportlab.platypus import Image
            logo_path = election.logo.path
            logo = Image(logo_path, width=1.5*inch, height=0.75*inch)
            logo.hAlign = 'CENTER'
            elements.append(logo)
            elements.append(Spacer(1, 12))
        except:
            pass  # Skip logo if can't load
    
    elements.append(Paragraph(f'Election Results: {election.title}', styles['Title']))
    elements.append(Spacer(1, 12))
    from collections import defaultdict
    position_candidates = defaultdict(list)
    for candidate in election.candidates.select_related('position').all():
        position_candidates[candidate.position].append(candidate)
    for position in sorted(position_candidates.keys(), key=lambda p: p.id):
        candidates = position_candidates[position]
        total = sum(c.vote_set.count() for c in candidates)
        elements.append(Paragraph(position.name, styles['Heading2']))
        data = [['Candidate', 'Rank', 'Force #', 'Votes', 'Percentage']]
        for c in candidates:
            votes = c.vote_set.count()
            pct = f'{(votes / total * 100):.1f}%' if total > 0 else '0.0%'
            data.append([c.name, c.get_rank_display(), c.force_number, str(votes), pct])
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="results_{election_id}.pdf"'
    response.write(pdf)
    return response

@login_required
@user_passes_test(is_admin)
def export_results_docx(request, election_id):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    election = get_object_or_404(Election.objects.prefetch_related('positions'), pk=election_id)
    create_audit_log(request.user, AuditLog.ACTION_EXPORT_RESULTS,
        f'Exported results for election: {election.title} (ID: {election.id}) as DOCX', request, 'Election', election.id)
    doc = Document()
    # Add logo if available
    if election.logo:
        try:
            logo_path = election.logo.path
            doc.add_picture(logo_path, width=Inches(2.5))
        except:
            pass
    
    title = doc.add_heading(f'Election Results: {election.title}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from collections import defaultdict
    position_candidates = defaultdict(list)
    for candidate in election.candidates.select_related('position').all():
        position_candidates[candidate.position].append(candidate)
    for position in sorted(position_candidates.keys(), key=lambda p: p.id):
        candidates = position_candidates[position]
        total = sum(c.vote_set.count() for c in candidates)
        doc.add_heading(position.name, level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Candidate', 'Rank', 'Force #', 'Votes', 'Percentage']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        for c in candidates:
            votes = c.vote_set.count()
            pct = f'{(votes / total * 100):.1f}%' if total > 0 else '0.0%'
            row_cells = table.add_row().cells
            row_cells[0].text = c.name
            row_cells[1].text = c.get_rank_display()
            row_cells[2].text = c.force_number
            row_cells[3].text = str(votes)
            row_cells[4].text = pct
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="results_{election_id}.docx"'
    return response

def _normalize_bulk_key(value):
    return str(value).strip().lower() if value is not None else ''

def _normalize_bulk_value(value):
    return str(value).strip() if value is not None else ''

def _load_bulk_voter_rows(uploaded_file, extension):
    rows = []
    if extension == 'csv':
        decoded = uploaded_file.read().decode('utf-8-sig')
        reader = csv.DictReader(io.StringIO(decoded))
        for row in reader:
            rows.append(row)
    elif extension in ('xlsx', 'xls'):
        try:
            import openpyxl
        except ImportError:
            raise ImportError('openpyxl is required for Excel support. Install: pip install openpyxl')
        wb = openpyxl.load_workbook(uploaded_file)
        ws = wb.active
        headers = [str(cell.value).strip() if cell.value else '' for cell in ws[1]]
        for row in ws.iter_rows(min_row=2, values_only=True):
            row_dict = {}
            for i, value in enumerate(row):
                if i < len(headers) and headers[i]:
                    row_dict[headers[i]] = str(value).strip() if value is not None else ''
            rows.append(row_dict)
    return rows

@login_required
@user_passes_test(is_admin)
def bulk_register_voters(request):
    created_voters = []
    errors = []
    if request.method == 'POST':
        form = BulkVoterUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['file']
            ext = uploaded_file.name.lower().split('.')[-1]
            try:
                rows = _load_bulk_voter_rows(uploaded_file, ext)
            except Exception as e:
                errors.append(f'Error reading file: {str(e)}')
                rows = []
            if rows:
                valid_ranks = {r[0] for r in PoliceUser._meta.get_field('rank').choices}
                for i, row in enumerate(rows, start=1):
                    try:
                        row_data = {_normalize_bulk_key(k): _normalize_bulk_value(v) for k, v in row.items() if k}
                        force_number = row_data.get('force_number', '')
                        first_name = row_data.get('first_name', '')
                        last_name = row_data.get('last_name', '')
                        email = row_data.get('email', '')
                        rank = row_data.get('rank', '').upper()
                        station = row_data.get('station', '')
                        phone = row_data.get('phone', '')
                        is_active_voter_str = row_data.get('is_active_voter', 'true')
                        if not force_number:
                            errors.append(f'Row {i}: force_number is required')
                            continue
                        try:
                            force_number = int(force_number)
                        except ValueError:
                            errors.append(f'Row {i}: force_number must be an integer')
                            continue
                        if PoliceUser.objects.filter(force_number=force_number).exists():
                            errors.append(f'Row {i}: Force number {force_number} already exists')
                            continue
                        if not first_name or not last_name:
                            errors.append(f'Row {i}: first_name and last_name are required')
                            continue
                        if not email:
                            errors.append(f'Row {i}: email is required')
                            continue
                        if not rank:
                            errors.append(f'Row {i}: rank is required')
                            continue
                        if rank not in valid_ranks:
                            errors.append(f'Row {i}: Invalid rank "{rank}". Valid: {", ".join(sorted(valid_ranks))}')
                            continue
                        if not station:
                            errors.append(f'Row {i}: station is required')
                            continue
                        auto_password = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(10))
                        username = str(force_number)
                        voter = PoliceUser.objects.create(
                            username=username, first_name=first_name, last_name=last_name,
                            email=email, force_number=force_number, rank=rank, station=station,
                            phone=phone, role='VOTER',
                            is_active_voter=str(is_active_voter_str).lower() in ('true', '1', 'yes', 'on'),
                            is_active=True, is_staff=False, must_change_password=True
                        )
                        voter.set_password(auto_password)
                        voter.save()
                        created_voters.append({
                            'row': i, 'username': username, 'force_number': force_number,
                            'full_name': voter.get_full_name(), 'email': email, 'password': auto_password
                        })
                    except Exception as e:
                        errors.append(f'Row {i}: {str(e)}')
                if created_voters:
                    create_audit_log(request.user, AuditLog.ACTION_VOTER_BULK_CREATE,
                        f'Bulk registered {len(created_voters)} voters', request)
                    if not errors:
                        messages.success(request, f'Successfully registered {len(created_voters)} voters!')
                elif not errors:
                    errors.append('No valid voter data found in the file.')
        else:
            errors.append('Invalid form submission.')
    else:
        form = BulkVoterUploadForm()
    return render(request, 'polls/bulk_register_voters.html', {
        'form': form, 'errors': errors, 'created_voters': created_voters
    })

@login_required
@user_passes_test(is_admin)
def audit_log(request):
    logs = AuditLog.objects.all().select_related('user').order_by('-timestamp')[:500]
    return render(request, 'polls/audit_log.html', {'logs': logs})

@login_required
@user_passes_test(is_admin)
def admin_election_voters(request, election_id):
    election = get_object_or_404(Election.objects.prefetch_related('positions'), pk=election_id)
    registered_voter_ids = set(ElectionRegistration.objects.filter(
        election=election
    ).values_list('voter_id', flat=True))
    
    registered_voters = PoliceUser.objects.filter(
        id__in=registered_voter_ids, role='VOTER'
    ).order_by('force_number')
    
    eligible_voters = []
    for voter in PoliceUser.objects.filter(role='VOTER', is_active=True).order_by('force_number'):
        if election.is_voter_eligible(voter) and voter.id not in registered_voter_ids:
            eligible_voters.append(voter)
    
    return render(request, 'polls/admin_election_voters.html', {
        'election': election,
        'registered_voters': registered_voters,
        'eligible_voters': eligible_voters,
        'registered_count': registered_voters.count(),
        'eligible_count': len(eligible_voters),
    })

@login_required
@user_passes_test(is_admin)
def admin_register_voter_to_election(request, election_id, voter_id):
    election = get_object_or_404(Election, pk=election_id)
    voter = get_object_or_404(PoliceUser, pk=voter_id, role='VOTER')
    
    if not election.is_voter_eligible(voter):
        messages.error(request, f'Voter {voter.get_full_name()} is not eligible for this election.')
        return redirect('polls:admin_election_voters', election_id=election_id)
    
    if ElectionRegistration.objects.filter(voter=voter, election=election).exists():
        messages.info(request, f'Voter {voter.get_full_name()} is already registered for this election.')
        return redirect('polls:admin_election_voters', election_id=election_id)

    ElectionRegistration.objects.create(voter=voter, election=election, registered_by=request.user)
    create_audit_log(
        request.user, AuditLog.ACTION_VOTER_UPDATE,
        f'Registered voter {voter.username} for election: {election.title}', request,
        target_model='Election', target_id=election.id
    )
    
    # Attempt to send election invitation email
    email_sent = send_election_invitation_email(voter, election)
    if email_sent:
        messages.success(request, f'Voter {voter.get_full_name()} registered for "{election.title}". Invitation sent to {voter.email}.')
    else:
        messages.success(request, f'Voter {voter.get_full_name()} registered for "{election.title}".')
        if voter.email:
            messages.warning(request, f'Could not send invitation email to {voter.email}.')
    
    return redirect('polls:admin_election_voters', election_id=election_id)

@login_required
@user_passes_test(is_admin)
def admin_unregister_voter_from_election(request, election_id, voter_id):
    election = get_object_or_404(Election, pk=election_id)
    voter = get_object_or_404(PoliceUser, pk=voter_id, role='VOTER')
    
    registration = ElectionRegistration.objects.filter(voter=voter, election=election).first()
    if not registration:
        messages.error(request, f'Voter {voter.get_full_name()} is not registered for this election.')
        return redirect('polls:admin_election_voters', election_id=election_id)
    
    if Vote.objects.filter(voter=voter, election=election).exists():
        messages.error(request, f'Cannot unregister {voter.get_full_name()} — they have already voted in this election.')
        return redirect('polls:admin_election_voters', election_id=election_id)
    
    registration.delete()
    create_audit_log(
        request.user, AuditLog.ACTION_VOTER_UPDATE,
        f'Unregistered voter {voter.username} from election: {election.title}', request,
        target_model='Election', target_id=election.id
    )
    messages.success(request, f'Voter {voter.get_full_name()} unregistered from "{election.title}".')
    return redirect('polls:admin_election_voters', election_id=election_id)

@login_required
@user_passes_test(is_admin)
def send_election_invitation(request, election_id, voter_id):
    """Send election invitation email to a voter."""
    election = get_object_or_404(Election, pk=election_id)
    voter = get_object_or_404(PoliceUser, pk=voter_id, role='VOTER')
    
    if not ElectionRegistration.objects.filter(voter=voter, election=election).exists():
        messages.error(request, f'Voter {voter.get_full_name()} is not registered for this election.')
        return redirect('polls:admin_election_voters', election_id=election_id)
    
    email_sent = send_election_invitation_email(voter, election)
    if email_sent:
        messages.success(request, f'Invitation email sent to {voter.email}')
    else:
        if voter.email:
            messages.error(request, f'Failed to send invitation email to {voter.email}.')
        else:
            messages.error(request, f'Voter has no email address.')
    
    return redirect('polls:admin_election_voters', election_id=election_id)

@login_required
@user_passes_test(is_superadmin)
def admin_reset_voter_votes(request, voter_id):
    voter = get_object_or_404(PoliceUser, pk=voter_id, role='VOTER')
    if request.method == 'POST':
        deleted_count, _ = Vote.objects.filter(voter=voter).delete()
        create_audit_log(
            request.user, AuditLog.ACTION_VOTER_UPDATE,
            f'Reset all votes for voter: {voter.username} (Force #{voter.force_number}) — {deleted_count} vote(s) deleted',
            request, 'PoliceUser', voter.id
        )
        messages.success(request, f'All votes reset for {voter.get_full_name()}. Deleted {deleted_count} vote(s).')
        return redirect('polls:admin_voters')
    return redirect('polls:admin_voters')

@login_required
@user_passes_test(is_superadmin)
def admin_reset_voter_election_votes(request, voter_id, election_id):
    voter = get_object_or_404(PoliceUser, pk=voter_id, role='VOTER')
    election = get_object_or_404(Election, pk=election_id)
    if request.method == 'POST':
        deleted_count, _ = Vote.objects.filter(voter=voter, election=election).delete()
        create_audit_log(
            request.user, AuditLog.ACTION_VOTER_UPDATE,
            f'Reset votes for voter: {voter.username} (Force #{voter.force_number}) in election "{election.title}" — {deleted_count} vote(s) deleted',
            request, 'Election', election.id
        )
        messages.success(request, f'Votes reset for {voter.get_full_name()} in "{election.title}". Deleted {deleted_count} vote(s).')
        return redirect('polls:admin_election_voters', election_id=election_id)
    return redirect('polls:admin_election_voters', election_id=election_id)
